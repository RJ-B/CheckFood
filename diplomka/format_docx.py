"""
Formátování diplomové práce dle Výnosu děkanky FIM č. 9/2025
+ úvodní strany dle vzoru z bakalářské práce autora.
+ klikatelné křížové odkazy a citace
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import re
import os
import copy

os.chdir(r"c:\Users\A\Desktop\WEBY\CheckFood\diplomka")

BOOKMARK_ID = [0]  # mutable counter

# Czech single-letter prepositions and conjunctions that must not be at end of line
NBSP = '\u00A0'
NBSP_PATTERN = re.compile(r'(?<= )([aikosuvzAIKOSUVZ]) ', re.UNICODE)

def fix_nbsp(text):
    """Replace normal space after single-letter Czech prepositions/conjunctions with non-breaking space."""
    return NBSP_PATTERN.sub(lambda m: m.group(1) + NBSP, text)

# =============================================================================
# HELPERS
# =============================================================================

def next_bookmark_id():
    BOOKMARK_ID[0] += 1
    return BOOKMARK_ID[0]


def set_font(run, name, size, bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._r
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def add_paragraph(doc, text, font_name="Cambria", font_size=12, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0,
                  space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(fix_nbsp(text))
    set_font(run, font_name, font_size, bold)
    return p


def add_page_number_to_footer(section, font_name="Cambria", font_size=11):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    for fld in [("begin", None), (None, " PAGE "), ("end", None)]:
        run = fp.add_run()
        if fld[0]:
            run._r.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="{fld[0]}"/>'))
        if fld[1]:
            run._r.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve">{fld[1]}</w:instrText>'))
        set_font(run, font_name, font_size)


def suppress_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()


def set_section_margins(section):
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    bid = next_bookmark_id()
    p_elem = paragraph._p
    bs = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{bookmark_name}"/>')
    be = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>')
    p_elem.append(bs)
    p_elem.append(be)


def make_hyperlink_run(paragraph, anchor, text, font_name="Cambria",
                       font_size=12, bold=False):
    """Create an internal hyperlink run inside a paragraph."""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    hyperlink = etree.SubElement(paragraph._p, qn('w:hyperlink'))
    hyperlink.set(qn('w:anchor'), anchor)

    run_elem = etree.SubElement(hyperlink, qn('w:r'))
    rPr = etree.SubElement(run_elem, qn('w:rPr'))

    # Font
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(int(font_size * 2)))
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), str(int(font_size * 2)))

    if bold:
        b = etree.SubElement(rPr, qn('w:b'))

    # Black, no underline
    color = etree.SubElement(rPr, qn('w:color'))
    color.set(qn('w:val'), '000000')

    t = etree.SubElement(run_elem, qn('w:t'))
    t.set(qn('xml:space'), 'preserve')
    t.text = text


def extract_heading_number(text):
    """Extract heading number like '1.3.2' from heading text like '1.3.2 Zabezpečení'."""
    m = re.match(r'^(\d+(?:\.\d+)*)\s', text.strip())
    if m:
        return m.group(1)
    return None


def heading_to_bookmark(num):
    """Convert heading number to bookmark name: '1.3.2' -> '_ref_1_3_2'."""
    return '_ref_' + num.replace('.', '_')


def citation_bookmark(num):
    """Bookmark name for citation: 3 -> '_cite_3'."""
    return f'_cite_{num}'


# Patterns for cross-references
# Matches: kapitole 1.3, sekci 1.3.2, kapitole 2, sekcí 1.2.1
XREF_PATTERN = re.compile(
    r'(kapitol[eěy]|kapitoly|kapitolu)\s+(\d+(?:\.\d+)*)')

# Matches: [1], [2], [15], [25] [27]
CITE_PATTERN = re.compile(r'\[(\d+)\]')


def process_text_with_links(paragraph, text, heading_bookmarks, font_name="Cambria",
                            font_size=12, bold=False, is_bibliography=False,
                            bib_number=None):
    """Process text, replacing cross-refs and citations with hyperlinks."""

    if is_bibliography and bib_number is not None:
        # Add bookmark to this bibliography entry
        add_bookmark(paragraph, citation_bookmark(bib_number))
        run = paragraph.add_run(fix_nbsp(text))
        set_font(run, font_name, font_size, bold)
        return

    # Combined pattern: find both xrefs and citations
    parts = []
    last_end = 0

    # Find all matches of both patterns
    matches = []
    for m in XREF_PATTERN.finditer(text):
        ref_num = m.group(2)
        bm = heading_to_bookmark(ref_num)
        if bm.replace('_ref_', '').replace('_', '.') in heading_bookmarks:
            matches.append((m.start(), m.end(), 'xref', m))
    for m in CITE_PATTERN.finditer(text):
        matches.append((m.start(), m.end(), 'cite', m))

    matches.sort(key=lambda x: x[0])

    for start, end, mtype, m in matches:
        # Add text before match
        if start > last_end:
            run = paragraph.add_run(fix_nbsp(text[last_end:start]))
            set_font(run, font_name, font_size, bold)

        if mtype == 'xref':
            ref_num = m.group(2)
            prefix = m.group(1)
            bm = heading_to_bookmark(ref_num)
            # Add prefix as normal text
            run = paragraph.add_run(prefix + ' ')
            set_font(run, font_name, font_size, bold)
            # Add number as hyperlink
            make_hyperlink_run(paragraph, bm, ref_num, font_name, font_size, bold)
        elif mtype == 'cite':
            cite_num = m.group(1)
            bm = citation_bookmark(cite_num)
            # Add [N] as hyperlink
            make_hyperlink_run(paragraph, bm, f'[{cite_num}]', font_name, font_size, bold)

        last_end = end

    # Add remaining text after last match, or full text if no matches
    if matches:
        if last_end < len(text):
            run = paragraph.add_run(fix_nbsp(text[last_end:]))
            set_font(run, font_name, font_size, bold)
    else:
        run = paragraph.add_run(fix_nbsp(text))
        set_font(run, font_name, font_size, bold)


# =============================================================================
# PHASE 1: Parse source document, collect heading bookmarks
# =============================================================================
src = Document("diplomova_prace_NOVA.docx")

heading_bookmarks = set()  # set of heading numbers like {'1', '1.1', '1.3.2', ...}

def is_heading(p):
    return "heading" in (p.style.name.lower() if p.style else "")

def get_heading_level(p):
    style = p.style.name if p.style else ""
    m = re.search(r'(\d+)', style)
    return int(m.group(1)) if m and "heading" in style.lower() else 0

# First pass: collect all heading numbers
for p in src.paragraphs:
    if is_heading(p):
        num = extract_heading_number(p.text.strip())
        if num:
            heading_bookmarks.add(num)

# Also add unnumbered headings as bookmarks
special_headings = {
    'Úvod': '_ref_uvod',
    'Teoretická část': '_ref_teoreticka',
    'Praktická část': '_ref_prakticka',
    'Závěr': '_ref_zaver',
    'Seznam použité literatury': '_ref_literatura',
}

print(f"Found {len(heading_bookmarks)} numbered headings: {sorted(heading_bookmarks)}")

# =============================================================================
# PHASE 2: Create the document
# =============================================================================
doc = Document()
set_section_margins(doc.sections[0])
suppress_page_number(doc.sections[0])

# --- Nastavení TOC stylů s odsazením pro úrovně ---
for level, indent_cm in [(1, 0), (2, 1.0), (3, 2.0)]:
    style_name = f'toc {level}'
    try:
        toc_style = doc.styles[style_name]
    except KeyError:
        toc_style = doc.styles.add_style(style_name, 1)  # 1 = paragraph style
    pf = toc_style.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.space_before = Pt(2 if level > 1 else 4)
    pf.space_after = Pt(2 if level > 1 else 4)
    pf.line_spacing = 1.15
    toc_style.font.name = "Cambria"
    toc_style.font.size = Pt(12 if level == 1 else 11)
    if level == 1:
        toc_style.font.bold = True
    else:
        toc_style.font.bold = False

# --- Nastavení Heading stylů ---
for level in [1, 2, 3]:
    style_name = f'Heading {level}'
    try:
        h_style = doc.styles[style_name]
    except KeyError:
        continue
    h_style.font.name = "Verdana"
    h_style.font.size = Pt(14)
    h_style.font.bold = True
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    pf = h_style.paragraph_format
    pf.keep_with_next = True

# =============================================================================
# TITULNÍ STRANA — tabulka 3 řádky x 1 sloupec (nahoře, střed, dole)
# =============================================================================
# Výška stránky bez okrajů: 29.7 - 2 - 2 = 25.7 cm
# Rozdělíme: nahoře ~6cm, střed ~10cm, dole ~9.7cm
title_table = doc.add_table(rows=3, cols=1)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Odstranit ohraničení tabulky
for row in title_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            tc.insert(0, tcPr)
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>')
        tcPr.append(borders)

# Nastavení výšky řádků
from docx.shared import Cm as CmShared
title_table.rows[0].height = Cm(6)
title_table.rows[1].height = Cm(10)
title_table.rows[2].height = Cm(9)

# --- NAHOŘE: Univerzita, Fakulta, Katedra ---
cell_top = title_table.cell(0, 0)
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
cell_top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

p = cell_top.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.clear()
run = p.add_run("Univerzita Hradec Králové")
set_font(run, "Cambria", 16, bold=True)

p2 = cell_top.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Fakulta informatiky a managementu")
set_font(run2, "Cambria", 14, bold=True)

p3 = cell_top.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Katedra informatiky a kvantitativních metod")
set_font(run3, "Cambria", 12, bold=False)

# --- STŘED: Název práce + typ ---
cell_mid = title_table.cell(1, 0)
cell_mid.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

p = cell_mid.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.clear()
run = p.add_run("Vývoj aplikace pro správu rezervací a objednávek v gastronomii")
set_font(run, "Cambria", 18, bold=True)

p2 = cell_mid.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(12)
run2 = p2.add_run("Diplomová práce")
set_font(run2, "Cambria", 14, bold=False)

# --- DOLE: Autor, obor, vedoucí, místo + datum ---
cell_bot = title_table.cell(2, 0)
cell_bot.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

p = cell_bot.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.clear()
p.paragraph_format.line_spacing = 1.15
run = p.add_run("Autor: Bc. Rostislav Jirák")
set_font(run, "Cambria", 12)

p2 = cell_bot.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.line_spacing = 1.15
run2 = p2.add_run("Studijní obor: Aplikovaná informatika")
set_font(run2, "Cambria", 12)

p3 = cell_bot.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.paragraph_format.line_spacing = 1.15
p3.paragraph_format.space_after = Pt(36)
run3 = p3.add_run("Vedoucí práce: doc. Mgr. Tomáš Kozel, Ph.D.")
set_font(run3, "Cambria", 12)

# Místo a datum - na dvou řádcích, centrovány
p4 = cell_bot.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.line_spacing = 1.15
p4.paragraph_format.space_after = Pt(0)
run4 = p4.add_run("Hradec Králové")
set_font(run4, "Cambria", 12)

p5 = cell_bot.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p5.paragraph_format.line_spacing = 1.15
run5 = p5.add_run("srpen 2026")
set_font(run5, "Cambria", 12)

# =============================================================================
# Helper: full-page table with content at bottom
# =============================================================================
def make_bottom_aligned_page(doc, build_content_fn):
    """Create a page where content is pushed to the bottom using a 1x1 borderless table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Full page height minus margins
    tbl.rows[0].height = Cm(25.7)
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    # Remove borders
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tcPr)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)
    build_content_fn(cell)

# =============================================================================
# PROHLÁŠENÍ (zarovnáno dole)
# =============================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
set_section_margins(doc.sections[-1])
suppress_page_number(doc.sections[-1])

def build_prohlaseni(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    run = p.add_run("Prohlášení")
    set_font(run, "Verdana", 14, bold=True)
    p.paragraph_format.space_after = Pt(12)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(fix_nbsp(
        'Prohlašuji, že jsem diplomovou práci na téma \u201eVývoj aplikace pro správu '
        'rezervací a objednávek v gastronomii\u201c vypracoval samostatně pod vedením '
        'vedoucího diplomové práce a s použitím odborné literatury a dalších '
        'informačních zdrojů, které jsou citovány v práci a uvedeny v seznamu '
        'použité literatury.'))
    set_font(run2, "Cambria", 12)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.space_after = Pt(24)
    run3 = p3.add_run(fix_nbsp(
        "Jako autor této diplomové práce dále prohlašuji, že v souvislosti "
        "s jejím vytvořením jsem neporušil autorská práva třetích osob."))
    set_font(run3, "Cambria", 12)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.paragraph_format.space_after = Pt(0)
    run4 = p4.add_run("V Hradci Králové dne")
    set_font(run4, "Cambria", 12)
    run4tab = p4.add_run("\t")
    set_font(run4tab, "Cambria", 12)
    run4sig = p4.add_run("podpis")
    set_font(run4sig, "Cambria", 12)
    # Pravý tab stop pro podpis
    pPr4 = p4._p.find(qn('w:pPr'))
    if pPr4 is None:
        pPr4 = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        p4._p.insert(0, pPr4)
    tabs4 = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9072"/></w:tabs>')
    pPr4.append(tabs4)

make_bottom_aligned_page(doc, build_prohlaseni)

# =============================================================================
# PODĚKOVÁNÍ (zarovnáno dole)
# =============================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
set_section_margins(doc.sections[-1])
suppress_page_number(doc.sections[-1])

def build_podekovani(cell):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    run = p.add_run("Poděkování")
    set_font(run, "Verdana", 14, bold=True)
    p.paragraph_format.space_after = Pt(12)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.line_spacing = 1.5
    run2 = p2.add_run(fix_nbsp(
        "Děkuji vedoucímu diplomové práce panu doc. Mgr. Tomáši Kozlovi, Ph.D. "
        "za odborné vedení, cenné rady a připomínky při zpracování této práce."))
    set_font(run2, "Cambria", 12)

make_bottom_aligned_page(doc, build_podekovani)

# =============================================================================
# ANOTACE
# =============================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
set_section_margins(doc.sections[-1])
suppress_page_number(doc.sections[-1])

add_paragraph(doc, "Anotace", font_name="Verdana", font_size=14,
              bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
add_paragraph(doc, "Diplomová práce se zabývá návrhem, implementací a testováním komplexního "
              "systému pro správu rezervací a objednávek v gastronomii. Na základě analýzy "
              "existujících řešení na trhu byla identifikována absence komplexního systému, "
              "který by na českém trhu kombinoval vyhledávání restaurací na mapě, interaktivní "
              "výběr stolu prostřednictvím 360° panoramatického zobrazení interiéru, správu "
              "rezervací a objednávek z mobilní aplikace a administrační nástroje pro personál "
              "a vlastníky restaurací. Výsledný systém CheckFood tuto mezeru zaplňuje. "
              "Serverová část je implementována ve frameworku Spring Boot s jazykem Java 21 "
              "a databází PostgreSQL s rozšířením PostGIS. Klientská mobilní aplikace je "
              "vyvinutá ve frameworku Flutter. Součástí systému je podpůrná mikroslužba "
              "v jazyce Python pro zpracování panoramatických snímků. Systém je nasazen "
              "v cloudovém prostředí s automatizovaným procesem sestavení, testování a nasazení.",
              font_size=12, space_after=12)
add_paragraph(doc, "Klíčová slova: Flutter, Spring Boot, PostgreSQL, rezervační systém, mobilní aplikace",
              font_size=12, bold=True, space_after=24)

add_paragraph(doc, "Annotation", font_name="Verdana", font_size=14,
              bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
add_paragraph(doc, "This diploma thesis deals with the design, implementation and testing of "
              "a comprehensive system for reservation and order management in gastronomy. "
              "Based on an analysis of existing solutions on the market, the absence of "
              "a comprehensive system was identified that would combine restaurant search on "
              "a map, interactive table selection through a 360° panoramic interior view, "
              "reservation and order management from a mobile application, and administrative "
              "tools for staff and restaurant owners on the Czech market. The resulting "
              "CheckFood system fills this gap. The server side is implemented in the "
              "Spring Boot framework with Java 21 and a PostgreSQL database with the PostGIS "
              "extension. The client mobile application is developed in the Flutter framework. "
              "The system includes a supporting microservice in Python for processing "
              "panoramic images. The system is deployed in a cloud environment with an "
              "automated build, test and deployment process.",
              font_size=12, space_after=12)
add_paragraph(doc, "Keywords: Flutter, Spring Boot, PostgreSQL, reservation system, mobile application",
              font_size=12, bold=True, space_after=0)

# =============================================================================
# OBSAH + VLASTNÍ TEXT (arabic page numbers from 1)
# =============================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sect_main = doc.sections[-1]
set_section_margins(sect_main)
add_page_number_to_footer(sect_main)

sectPr = sect_main._sectPr
pgNumType = sectPr.find(qn('w:pgNumType'))
if pgNumType is None:
    pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
    sectPr.append(pgNumType)
pgNumType.set(qn('w:start'), '1')

# OBSAH - TOC field
add_paragraph(doc, "OBSAH", font_name="Verdana", font_size=14,
              bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

p_toc = doc.add_paragraph()
run1 = p_toc.add_run()
run1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
run2 = p_toc.add_run()
run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
run3 = p_toc.add_run()
run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
run4 = p_toc.add_run("Pravý klik → Aktualizovat pole → Aktualizovat celou tabulku")
set_font(run4, "Cambria", 11)
run5 = p_toc.add_run()
run5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

p_break = doc.add_paragraph()
p_break.add_run().add_break(WD_BREAK.PAGE)

# =============================================================================
# IMPORT TEXT with clickable cross-references and citations
# =============================================================================

skip_until_uvod = True
in_bibliography = False

for paragraph in src.paragraphs:
    text = paragraph.text.strip()

    # Skip until Úvod
    if skip_until_uvod:
        if is_heading(paragraph) and "vod" in text.lower():
            skip_until_uvod = False
        else:
            continue

    # Skip separators and MD title page remnants
    if text in ['---', '']:
        continue
    if text in ['Univerzita Hradec Králové', 'Fakulta informatiky a managementu',
                'Katedra informatiky a kvantitativních metod']:
        continue

    # Detect bibliography section
    if is_heading(paragraph) and 'seznam' in text.lower() and 'liter' in text.lower():
        in_bibliography = True

    if is_heading(paragraph):
        src_level = get_heading_level(paragraph)
        if src_level == 0:
            src_level = 1

        # Remap pandoc heading levels to Word heading levels:
        # MD # (H1) = section dividers (Teoretická/Praktická část) → Heading 1
        # MD ## (H2) = main chapters (Úvod, 1, 2, 3..., Závěr) → Heading 1
        # MD ### (H3) = subchapters (1.1, 2.3...) → Heading 2
        # MD #### (H4) = sub-subchapters (1.2.1, 2.3.4...) → Heading 3
        if src_level <= 2:
            doc_level = 1
        elif src_level == 3:
            doc_level = 2
        else:
            doc_level = 3

        style_name = f'Heading {doc_level}'
        p = doc.add_paragraph(style=style_name)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format

        if doc_level == 1:
            pf.page_break_before = True
            pf.space_before = Pt(24)
            pf.space_after = Pt(12)
        elif doc_level == 2:
            pf.space_before = Pt(18)
            pf.space_after = Pt(8)
        else:
            pf.space_before = Pt(14)
            pf.space_after = Pt(6)

        run = p.add_run(fix_nbsp(text))
        set_font(run, "Verdana", 14, bold=True)

        # Add bookmark for this heading
        hnum = extract_heading_number(text)
        if hnum:
            add_bookmark(p, heading_to_bookmark(hnum))
        elif text in special_headings:
            add_bookmark(p, special_headings[text])

    elif text.startswith("**Tabulka ") or text.startswith("*Obrázek ") or re.match(r'^Tabulka \d+:', text) or re.match(r'^Obrázek \d+:', text):
        clean = text.strip('*')

        # Check if there's an image file to insert
        IMAGE_MAP = {
            'Obrázek 1:': 'diagrams/BLoC to Widget State Flow-2026-04-06-210359.png',
            'Obrázek 2:': 'diagrams/BLoC to Widget State Flow-2026-04-06-210643.png',
        }
        img_path = None
        for key, path in IMAGE_MAP.items():
            if clean.startswith(key):
                full_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), path)
                if os.path.exists(full_path):
                    img_path = full_path
                break

        if img_path:
            # Insert image centered (caption added manually in Word via Insert Caption)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(12)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Cm(15))
        else:
            # No image file — add text placeholder for caption
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.5
            run = p.add_run(clean)
            set_font(run, "Cambria", 11, bold=True)

    elif '|' in text and text.startswith('|'):
        continue  # skip MD table rows

    else:
        # Normal paragraph with clickable references
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)

        # Check if this is a bibliography entry like [1] SENSOR TOWER...
        bib_match = re.match(r'^\[(\d+)\]\s', text)
        if in_bibliography and bib_match:
            bib_num = int(bib_match.group(1))
            add_bookmark(p, citation_bookmark(bib_num))
            run = p.add_run(text)
            set_font(run, "Cambria", 12)
        else:
            # Get full text from all runs
            full_text = ''.join(r.text for r in paragraph.runs) if paragraph.runs else text
            if not full_text.strip():
                full_text = text
            process_text_with_links(p, full_text, heading_bookmarks, "Cambria", 12)

# Copy tables from source
for src_table in src.tables:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(src_table.rows):
        for j, cell in enumerate(row.cells):
            dst_cell = table.cell(i, j)
            dst_cell.paragraphs[0].clear()

            for k, src_p in enumerate(cell.paragraphs):
                p = dst_cell.paragraphs[0] if k == 0 else dst_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.line_spacing = 1.15
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)

                for src_run in src_p.runs:
                    run = p.add_run(src_run.text)
                    set_font(run, "Cambria", 10, bold=(i == 0) or (src_run.bold or False))

# =============================================================================
# SAVE
# =============================================================================
doc.save("diplomova_prace_FORMATOVANA.docx")
print("Hotovo: diplomova_prace_FORMATOVANA.docx")
